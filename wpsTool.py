import os
import docx
import copy
import time
from docx.shared import Pt  # 设置像素、缩进等
from win32com import client as wc
from openpyxl import load_workbook

TitleDic = {"车辆识别代码": "", "申请人姓名": "", "申请人性质": "", "证件类型": "",
            "证件号": "", "申请人电话": 0, "汽车销售商": "", "车辆类型": "插电式",
            "充电设备类型": "220V交流", "充电设备性质": "私人", "充电设备安装区域": "",
            "充电设备安装街区": "", "充电设备停车位": "", "充电桩编号": "",
            "充电设施建设运营单位": "上海莘永电力工程有限公司", "投入运营时间": ""
            }


def get_all_doc(dirPath):
    """
    遍历当前所有目录，获得Word列表、Excel表、图片
    返回值为[word_list, excel_path, picture_list]
    """
    word_list = []
    picture_list = []
    excel_path = ''
    tf = Transfer_format()

    for root_dir, sub_dir, files in os.walk(r'' + dirPath):
        for file in files:
            if file.endswith('.doc') or file.endswith('.docx'):
                file_name = os.path.join(root_dir, file)
                if file.endswith('.doc'):
                    tf.word_suffix(file_name)
                    os.remove(file_name)
                    file_name += "x"
                word_list.append(file_name)

            if file.endswith('.xlsx'):  # and '模板' in file
                excel_path = os.path.join(root_dir, file)

            if file.endswith('.jpg') or file.endswith('.png'):
                picture_list.append(os.path.join(root_dir, file))

    return [word_list, excel_path, picture_list]


class ReadWrite_excel(object):
    """修改excel数据"""
    dayTime = time.strftime("%Y/%m/%d", time.localtime())

    def __init__(self, filename):
        self.currRow = 2
        self.readRow = 1
        self.filename = filename
        self.wb = load_workbook(self.filename)
        self.ws = self.wb.active  # 激活sheet
        self.getBlackRow()

    def write(self, row_n, col_n, value):
        """写入数据，如(2,3,"hello"),第二行第三列写入数据"hello\""""
        self.ws.cell(row_n, col_n, value)
        self.wb.save(self.filename)

    def loop_write(self):
        column = 2
        for value in copy.deepcopy(TitleDic).values():
            self.ws.cell(self.currRow, column, value)
            column += 1
        self.currRow += 1
        self.dic_restore()
        self.wb.save(self.filename)

    def getBlackRow(self):
        # for row in range(1, self.ws.max_row):
        #     if not self.ws.cell(row, 2).value:
        #         self.currRow = row
        #         break
        return

    def dic_restore(self):
        sour = {"车辆识别代码": "", "申请人姓名": "", "申请人性质": "", "证件类型": "",
                "证件号": 0, "申请人电话": 0, "汽车销售商": "", "车辆类型": "插电式",
                "充电设备类型": "220V交流", "充电设备性质": "私人", "充电设备安装区域": "",
                "充电设备安装街区": "", "充电设备停车位": "", "充电桩编号": 0,
                "充电设施建设运营单位": "上海莘永电力工程有限公司", "投入运营时间": ""
                }
        for k in sour.keys():
            TitleDic[k] = sour[k]
        TitleDic["投入运营时间"] = self.dayTime

    def loop_read(self,word_list, picture_list):
        excelCount_list = []
        for self.readRow in range(1, self.ws.max_row + 1):
            row_dic = {"申请人姓名": self.ws.cell(self.readRow, 1).value,
                       "申请人电话": self.ws.cell(self.readRow, 2).value,
                       "充电设备安装地址": self.ws.cell(self.readRow, 3).value,
                       "充电桩编码": self.ws.cell(self.readRow, 4).value,
                       "充电设备停车位": ""
                       }
            temp = row_dic["充电设备安装地址"]
            if temp.find("【") != -1 and temp.rfind("】") != -1:
                row_dic["充电设备停车位"] = temp[temp.find("【") + 1:temp.find("】")]
                self.write(self.readRow, 3, temp[:temp.find("【")]
                           + temp[temp.find("【") + 1:temp.find("】")]
                           + temp[temp.find("】") + 1:])
            excelCount_list.append(row_dic)

        return excelCount_list


class Transfer_format(object):

    def __init__(self):
        self.word = wc.gencache.EnsureDispatch('kwps.application')

    def word_suffix(self, file_name):
        if os.path.exists("{}x".format(file_name)):
            return
        doc = self.word.Documents.Open(file_name)
        doc.SaveAs("{}x".format(file_name), 12)
        doc.Close()

    def word2pdf(self, file_name):
        # if os.path.exists(file_name[:file_name.rfind("doc")] + "pdf"):
        #     return
        doc = self.word.Documents.Open(file_name)
        doc.SaveAs(file_name[:file_name.rfind("doc")] + "pdf", FileFormat=17)
        doc.Close()

    @staticmethod
    def ModifyWordName(fn, VIC):
        """Vehicle identification code，车辆识别代码"""
        if "-自查表" not in fn:  # os.path.join(root_dir, file)
            strpath = fn[:(fn.rfind("\\") + 1)] + str(VIC) + "-自查表.docx"
            print(strpath, VIC)
            if not os.path.exists(strpath):
                os.rename(fn, strpath)
                fn = strpath
        return fn

    def __del__(self):
        self.word.Quit()


def Read_DocData(fn):
    """读取doc文档中需要的内容"""

    if not fn:
        print("Word文档路径为空")
        return

    doc = docx.Document(fn)

    # 按段落读取
    for paragraph in doc.paragraphs:
        if "经销商全称：" in paragraph.text:
            t = paragraph.text
            TitleDic["汽车销售商"] = paragraph.text[t.find("上海"):t.find("公司") + 3]

    VIN = ''
    applicantName_ls = []
    # 按表格读取全部数据
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if '＊车架号' in cell.text:
                    # LC0CE4CB0N0045391-自查表
                    VIN = cell.text[5:].strip()
                    TitleDic["车辆识别代码"] = VIN
                    print(VIN)
                # if '＊申请人' in cell.text:
                #     applicantName_ls.append(cell)
                if '＊申请人性质' in cell.text:
                    TitleDic["申请人性质"] = '私人'
                if '＊申请人证件类型' in cell.text:
                    TitleDic["证件类型"] = '居民身份证'
                if '＊证件号码' in cell.text:
                    temp = cell.text[6:].strip()
                    # excel表格数据精度为15位，身份证为18
                    TitleDic["证件号"] = temp
                if '申请人电话' in cell.text:
                    temp = cell.text[6:].strip()
                    TitleDic["申请人电话"] = int(temp) if temp.isdigit() else temp
                if '＊充电设施编号' in cell.text:
                    temp = cell.text[8:].strip()
                    TitleDic["充电桩编号"] = temp
                if '＊充电设施安装地址' in cell.text:
                    CIA = cell.text
                    if CIA.find("市") == -1:
                        TitleDic["充电设备安装区域"] = cell.text[10:CIA.find("区") + 1].strip()
                    else:
                        TitleDic["充电设备安装区域"] = cell.text[CIA.find("市") + 1:CIA.find("区") + 1].strip()
                    TitleDic["充电设备安装街区"] = cell.text[CIA.find("区") + 1:CIA.find("\n")].strip()
                if '车位号' in cell.text:
                    TitleDic["充电设备停车位"] = cell.text[4:].strip()
        ap = table.cell(1, 0).text
        if ap[4] != "：":
            table.cell(1, 0).text = ""
            run = table.cell(1, 0).paragraphs[0].add_run("＊申请人：" + ap[5:].strip())
            run.font.name = '仿宋_GB2312'
            run.font.size = Pt(12)
        TitleDic["申请人姓名"] = ap[5:].strip()
        doc.save(fn)
        # if len(applicantName_ls):
        #     ap = applicantName_ls[0].text
        #     print(ap)
        #     if ap[4] != "：":
        #         ap = "＊申请人：" + ap[5:].strip()
        #         ap.font.size = Pt(18)
        #         ap.font.name = "黑体"
        #     TitleDic["申请人姓名"] = ap[5:].strip()
        #     doc.save(fn)


def Write_Doc(fn, picturePath, row_dic):
    """写入车位号并且插入图片"""
    if not fn:
        print(row_dic["申请人姓名"]+"的word路径为空")
        return

    doc = docx.Document(fn)
    table = doc.tables[0]
    if row_dic["充电设备停车位"]:
        table.cell(7, 0).text = ""
        run = table.cell(7, 0).paragraphs[0].add_run("车位号：" + row_dic["充电设备停车位"])
        run.font.name = '仿宋_GB2312'
        run.font.size = Pt(12)
        doc.save(fn)
    else:
        print(row_dic["申请人姓名"], "在表格中没加入【 】标志，故无法提取充电设备停车位")

    VIC = table.cell(4, 3).text[5:].strip()
    # print("车编号为：" + VIC)
    if not picturePath:
        print(row_dic["申请人姓名"] + "的图片路径为空")
        return

    from PIL import Image
    f = Image.open(picturePath)  # 你的图片文件
    f.save(picturePath)  # 替换掉你的图片文件
    f.close()
    doc.add_picture(picturePath, height=Pt(200), width=Pt(450))
    doc.save(fn)
    strpath = picturePath[:(picturePath.rfind("\\") + 1)] + str(VIC) + ".jpg"
    if not os.path.exists(strpath):
        os.rename(picturePath, strpath)
